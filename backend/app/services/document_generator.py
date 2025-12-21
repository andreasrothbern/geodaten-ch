"""
Materialbewirtschaftung Document Generator

Generiert professionelle Word/PDF-Dokumente für die Gruppenleiter-Schulung
basierend auf realen Gebäudedaten aus der Schweiz.

Struktur (nach GL 2025 Vorlage):
1. Baustellenbeschrieb
2. Ausmass (NPK 114)
3. Materialauszug (Layher Blitz 70)
4. Personalbedarf
5. Dokumentation Baustelle
6. Reflexion
7. Anhang (SVG-Visualisierungen)
"""

from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any
from datetime import datetime
from io import BytesIO
import math

# Document generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class BuildingData:
    """Gebäudedaten für die Dokumentgenerierung"""
    address: str
    egid: Optional[int] = None

    # Gebäudemasse
    length_m: float = 10.0
    width_m: float = 10.0
    eave_height_m: float = 6.5
    ridge_height_m: Optional[float] = None
    floors: int = 2

    # GWR-Daten
    building_category: str = "Einfamilienhaus"
    construction_year: Optional[int] = None
    area_m2: Optional[float] = None

    # Dachform
    roof_type: str = "satteldach"  # satteldach, flachdach, walmdach
    roof_angle_deg: float = 35.0

    # Koordinaten
    lv95_e: Optional[float] = None
    lv95_n: Optional[float] = None


@dataclass
class ScaffoldRequirements:
    """Gerüstanforderungen"""
    scaffold_type: str = "Fassadengerüst, umlaufend (4 Seiten)"
    scaffold_system: str = "Layher Blitz 70 Stahl"
    load_class: int = 3  # 200 kg/m²
    width_class: str = "W09"  # 0.90 m Belagbreite
    facade_distance_m: float = 0.30
    has_gable_scaffold: bool = True
    has_roof_protection: bool = True


@dataclass
class NPK114Result:
    """NPK 114 Ausmass-Ergebnis"""
    facade_area_total_m2: float
    facade_a_c_m2: float  # Traufseiten
    facade_b_d_m2: float  # Giebelseiten
    corner_surcharge_m2: float

    # Einzelne Berechnungen
    calculations: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class MaterialList:
    """Materialliste"""
    items: List[Dict[str, Any]] = field(default_factory=list)
    total_weight_kg: float = 0.0

    # Nach Kategorien
    vertical_frames_kg: float = 0.0
    horizontal_frames_kg: float = 0.0
    decks_kg: float = 0.0
    diagonals_kg: float = 0.0
    base_plates_kg: float = 0.0
    anchoring_kg: float = 0.0
    accessories_kg: float = 0.0


@dataclass
class PersonnelRequirement:
    """Personalbedarf"""
    assembly_hours: float = 8.5
    disassembly_hours: float = 6.5
    team_size: int = 3

    @property
    def total_man_hours(self) -> float:
        return (self.assembly_hours + self.disassembly_hours) * self.team_size


class DocumentGenerator:
    """Generiert Materialbewirtschaftungs-Dokumente"""

    def __init__(self):
        self.company_name = "Lawil Gerüstbau AG"
        self.company_address = "Murtenstrasse 30, 3202 Frauenkappelen"
        self.company_phone = "031 920 00 30"
        self.company_email = "west@lawil.ch"
        self.company_web = "www.lawil.ch"

    def calculate_npk114(self, building: BuildingData, requirements: ScaffoldRequirements) -> NPK114Result:
        """Berechnet das Ausmass nach NPK 114 D/2012"""

        # NPK 114 Zuschläge
        LS = 1.0  # Stirnseitiger Abschluss (LF + LG = 0.30 + 0.70)
        height_surcharge = 1.0  # Höhenzuschlag über Arbeitshöhe

        calculations = []

        # Fassade A und C (Traufseiten - längere Seiten)
        L_trauf = max(building.length_m, building.width_m)
        LA_trauf = LS + L_trauf + LS
        HA_trauf = building.eave_height_m + height_surcharge
        area_trauf_single = LA_trauf * HA_trauf
        area_trauf_total = 2 * area_trauf_single

        calculations.append({
            "name": "Fassade A+C (Traufseiten)",
            "length_m": L_trauf,
            "LA_m": LA_trauf,
            "height_m": building.eave_height_m,
            "HA_m": HA_trauf,
            "area_single_m2": round(area_trauf_single, 2),
            "area_total_m2": round(area_trauf_total, 2),
            "count": 2
        })

        # Fassade B und D (Giebelseiten - kürzere Seiten)
        L_giebel = min(building.length_m, building.width_m)
        LA_giebel = LS + L_giebel + LS

        # Bei Satteldach: mittlere Höhe für Giebelseite
        if building.roof_type == "satteldach" and building.ridge_height_m:
            giebel_height = building.ridge_height_m - building.eave_height_m
            mean_height = building.eave_height_m + giebel_height * 0.5
        else:
            mean_height = building.eave_height_m

        HA_giebel = mean_height + height_surcharge
        area_giebel_single = LA_giebel * HA_giebel
        area_giebel_total = 2 * area_giebel_single

        calculations.append({
            "name": "Fassade B+D (Giebelseiten)",
            "length_m": L_giebel,
            "LA_m": LA_giebel,
            "height_m": mean_height,
            "HA_m": HA_giebel,
            "area_single_m2": round(area_giebel_single, 2),
            "area_total_m2": round(area_giebel_total, 2),
            "count": 2
        })

        # Eckzuschläge (4 Ecken)
        corner_area_single = LS * HA_trauf
        corner_area_total = 4 * corner_area_single

        calculations.append({
            "name": "Eckzuschläge",
            "area_single_m2": round(corner_area_single, 2),
            "area_total_m2": round(corner_area_total, 2),
            "count": 4
        })

        total_area = area_trauf_total + area_giebel_total + corner_area_total

        return NPK114Result(
            facade_area_total_m2=round(total_area, 2),
            facade_a_c_m2=round(area_trauf_total, 2),
            facade_b_d_m2=round(area_giebel_total, 2),
            corner_surcharge_m2=round(corner_area_total, 2),
            calculations=calculations
        )

    def estimate_material(self, npk_result: NPK114Result, building: BuildingData) -> MaterialList:
        """Schätzt den Materialbedarf basierend auf der Gerüstfläche"""

        area = npk_result.facade_area_total_m2
        perimeter = 2 * (building.length_m + building.width_m)
        height = building.eave_height_m

        # Anzahl Gerüstlagen (2m Stellrahmen)
        layers = math.ceil(height / 2.0)

        # Anzahl Felder pro Seite (ca. 3m Feldlänge)
        fields_long = math.ceil(max(building.length_m, building.width_m) / 3.07)
        fields_short = math.ceil(min(building.length_m, building.width_m) / 2.57)
        total_fields = 2 * fields_long + 2 * fields_short

        # Anzahl Ständer (Ecken + Zwischenständer)
        stands = total_fields + 4  # Felder + Ecken

        items = []

        # Vertikalrahmen / Stellrahmen
        stellrahmen_200 = stands * layers
        stellrahmen_100 = stands  # Ausgleich oben
        stellrahmen_050 = 8  # Reserve/Ausgleich

        items.extend([
            {"category": "Vertikalrahmen", "name": "Stellrahmen 2.00 m", "art_nr": "1201",
             "quantity": stellrahmen_200, "weight_kg": 18.5, "total_kg": stellrahmen_200 * 18.5},
            {"category": "Vertikalrahmen", "name": "Stellrahmen 1.00 m", "art_nr": "1202",
             "quantity": stellrahmen_100, "weight_kg": 12.0, "total_kg": stellrahmen_100 * 12.0},
            {"category": "Vertikalrahmen", "name": "Stellrahmen 0.50 m", "art_nr": "1203",
             "quantity": stellrahmen_050, "weight_kg": 8.5, "total_kg": stellrahmen_050 * 8.5},
        ])

        # Horizontalrahmen / Geländer
        gelaender_307 = 2 * fields_long * layers
        gelaender_257 = 2 * fields_short * layers
        gelaender_207 = int(total_fields * 0.3)  # Ausgleich
        stirngelaender = stands

        items.extend([
            {"category": "Horizontalrahmen", "name": "Doppelgeländer 3.07 m", "art_nr": "2301",
             "quantity": gelaender_307, "weight_kg": 10.5, "total_kg": gelaender_307 * 10.5},
            {"category": "Horizontalrahmen", "name": "Doppelgeländer 2.57 m", "art_nr": "2302",
             "quantity": gelaender_257, "weight_kg": 9.0, "total_kg": gelaender_257 * 9.0},
            {"category": "Horizontalrahmen", "name": "Doppelgeländer 2.07 m", "art_nr": "2303",
             "quantity": gelaender_207, "weight_kg": 7.5, "total_kg": gelaender_207 * 7.5},
            {"category": "Horizontalrahmen", "name": "Stirngeländer 0.73 m", "art_nr": "2401",
             "quantity": stirngelaender, "weight_kg": 4.0, "total_kg": stirngelaender * 4.0},
        ])

        # Beläge
        belag_307 = 2 * fields_long * layers
        belag_257 = 2 * fields_short * layers
        belag_207 = int(total_fields * 0.3)
        durchstieg = 4  # Pro Aufstieg

        items.extend([
            {"category": "Beläge", "name": "Robustboden 3.07 × 0.32 m", "art_nr": "3101",
             "quantity": belag_307, "weight_kg": 19.5, "total_kg": belag_307 * 19.5},
            {"category": "Beläge", "name": "Robustboden 2.57 × 0.32 m", "art_nr": "3102",
             "quantity": belag_257, "weight_kg": 16.5, "total_kg": belag_257 * 16.5},
            {"category": "Beläge", "name": "Robustboden 2.07 × 0.32 m", "art_nr": "3103",
             "quantity": belag_207, "weight_kg": 13.5, "total_kg": belag_207 * 13.5},
            {"category": "Beläge", "name": "Durchstiegsboden 3.07 × 0.64 m", "art_nr": "3201",
             "quantity": durchstieg, "weight_kg": 38.0, "total_kg": durchstieg * 38.0},
        ])

        # Diagonalen
        diag_307 = int(total_fields * 0.4)
        diag_257 = int(total_fields * 0.2)
        horiz = int(total_fields * 0.3)

        items.extend([
            {"category": "Diagonalen", "name": "Diagonale 3.07 m", "art_nr": "4101",
             "quantity": diag_307, "weight_kg": 5.5, "total_kg": diag_307 * 5.5},
            {"category": "Diagonalen", "name": "Diagonale 2.57 m", "art_nr": "4102",
             "quantity": diag_257, "weight_kg": 4.5, "total_kg": diag_257 * 4.5},
            {"category": "Diagonalen", "name": "Horizontalstrebe 3.07 m", "art_nr": "4201",
             "quantity": horiz, "weight_kg": 4.0, "total_kg": horiz * 4.0},
        ])

        # Fussplatten und Spindeln
        fussplatten = stands
        spindeln = stands

        items.extend([
            {"category": "Fussplatten", "name": "Fussplatte 150 × 150 mm", "art_nr": "5001",
             "quantity": fussplatten, "weight_kg": 2.5, "total_kg": fussplatten * 2.5},
            {"category": "Fussplatten", "name": "Fußspindel 0.40 m", "art_nr": "5101",
             "quantity": spindeln, "weight_kg": 3.0, "total_kg": spindeln * 3.0},
        ])

        # Verankerung
        geruesthalter = int(stands * 0.6)
        v_anker = 8
        ringoesen = int(stands * 0.8)

        items.extend([
            {"category": "Verankerung", "name": "Gerüsthalter kurz", "art_nr": "6001",
             "quantity": geruesthalter, "weight_kg": 1.5, "total_kg": geruesthalter * 1.5},
            {"category": "Verankerung", "name": "V-Anker", "art_nr": "6002",
             "quantity": v_anker, "weight_kg": 3.0, "total_kg": v_anker * 3.0},
            {"category": "Verankerung", "name": "Ringöse M12 mit Dübel", "art_nr": "6101",
             "quantity": ringoesen, "weight_kg": 0.3, "total_kg": ringoesen * 0.3},
        ])

        # Konsolen und Zubehör
        konsolen = 16
        bordbretter_307 = int(belag_307 * 0.4)
        bordbretter_257 = int(belag_257 * 0.4)
        leitern = 4

        items.extend([
            {"category": "Zubehör", "name": "Innenkonsole 0.36 m", "art_nr": "7001",
             "quantity": konsolen, "weight_kg": 6.5, "total_kg": konsolen * 6.5},
            {"category": "Zubehör", "name": "Bordbretter 3.07 m", "art_nr": "7101",
             "quantity": bordbretter_307, "weight_kg": 4.5, "total_kg": bordbretter_307 * 4.5},
            {"category": "Zubehör", "name": "Bordbretter 2.57 m", "art_nr": "7102",
             "quantity": bordbretter_257, "weight_kg": 3.8, "total_kg": bordbretter_257 * 3.8},
            {"category": "Zubehör", "name": "Leiter 2.00 m (Aufstieg)", "art_nr": "7201",
             "quantity": leitern, "weight_kg": 8.0, "total_kg": leitern * 8.0},
        ])

        # Gewichte nach Kategorie berechnen
        result = MaterialList(items=items)

        for item in items:
            category = item["category"]
            weight = item["total_kg"]
            result.total_weight_kg += weight

            if category == "Vertikalrahmen":
                result.vertical_frames_kg += weight
            elif category == "Horizontalrahmen":
                result.horizontal_frames_kg += weight
            elif category == "Beläge":
                result.decks_kg += weight
            elif category == "Diagonalen":
                result.diagonals_kg += weight
            elif category == "Fussplatten":
                result.base_plates_kg += weight
            elif category == "Verankerung":
                result.anchoring_kg += weight
            elif category == "Zubehör":
                result.accessories_kg += weight

        return result

    def calculate_personnel(self, npk_result: NPK114Result) -> PersonnelRequirement:
        """Berechnet den Personalbedarf basierend auf der Gerüstfläche"""

        area = npk_result.facade_area_total_m2

        # Erfahrungswerte: ca. 55 m²/h bei 3 Mann für Montage
        assembly_rate = 55.0  # m² pro Stunde bei 3 Mann

        # Grundzeit + flächenabhängige Zeit
        base_time = 2.5  # Abladen, Verankerung, Kontrolle
        assembly_time = area / assembly_rate
        total_assembly = base_time + assembly_time

        # Demontage ca. 20% schneller
        total_disassembly = total_assembly * 0.8

        return PersonnelRequirement(
            assembly_hours=round(total_assembly, 1),
            disassembly_hours=round(total_disassembly, 1),
            team_size=3
        )

    def generate_word_document(
        self,
        building: BuildingData,
        author_name: str = "Teilnehmer GL 2025",
        project_description: str = "Fassadensanierung",
        include_reflexion_template: bool = True
    ) -> bytes:
        """Generiert ein Word-Dokument (.docx) für die Materialbewirtschaftung"""

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen.")

        # Berechnungen durchführen
        requirements = ScaffoldRequirements()
        npk_result = self.calculate_npk114(building, requirements)
        material = self.estimate_material(npk_result, building)
        personnel = self.calculate_personnel(npk_result)

        # Dokument erstellen
        doc = Document()

        # Dokumenteigenschaften
        doc.core_properties.author = author_name
        doc.core_properties.title = f"Materialbewirtschaftung - {building.address}"

        # Styles anpassen
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # === DECKBLATT ===
        self._add_cover_page(doc, building, author_name, project_description)

        # === INHALTSVERZEICHNIS ===
        doc.add_page_break()
        self._add_table_of_contents(doc)

        # === 1. BAUSTELLENBESCHRIEB ===
        doc.add_page_break()
        self._add_section_1(doc, building, requirements, project_description)

        # === 2. AUSMASS ===
        doc.add_page_break()
        self._add_section_2(doc, building, npk_result)

        # === 3. MATERIALAUSZUG ===
        doc.add_page_break()
        self._add_section_3(doc, material)

        # === 4. PERSONALBEDARF ===
        doc.add_page_break()
        self._add_section_4(doc, npk_result, personnel)

        # === 5. DOKUMENTATION BAUSTELLE ===
        doc.add_page_break()
        self._add_section_5(doc, building, material)

        # === 6. REFLEXION ===
        if include_reflexion_template:
            doc.add_page_break()
            self._add_section_6(doc)

        # === 7. ANHANG ===
        doc.add_page_break()
        self._add_section_7(doc, building)

        # Als Bytes zurückgeben
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_cover_page(self, doc: Document, building: BuildingData, author: str, project: str):
        """Fügt das Deckblatt hinzu"""

        # Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("POLYBAU / POLYBAT")
        run.bold = True
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Gruppenleiter 2025")

        # Abstand
        for _ in range(3):
            doc.add_paragraph()

        # Titel
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Praxis-Umsetzung")
        run.bold = True
        run.font.size = Pt(24)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Materialbewirtschaftung")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 112, 192)  # Blau

        # Abstand
        for _ in range(2):
            doc.add_paragraph()

        # Untertitel
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Fassadengerüst {building.building_category}")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("inkl. Giebelgerüst und Dachfangschutz")
        run.italic = True

        # Abstand
        for _ in range(4):
            doc.add_paragraph()

        # Metadaten
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Verfasser:"
        table.cell(0, 0).paragraphs[0].runs[0].bold = True
        table.cell(0, 1).text = author

        table.cell(1, 0).text = "Baustelle:"
        table.cell(1, 0).paragraphs[0].runs[0].bold = True
        table.cell(1, 1).text = building.address

        table.cell(2, 0).text = "Datum:"
        table.cell(2, 0).paragraphs[0].runs[0].bold = True
        table.cell(2, 1).text = datetime.now().strftime("%B %Y")

    def _add_table_of_contents(self, doc: Document):
        """Fügt das Inhaltsverzeichnis hinzu"""

        h = doc.add_heading("Inhaltsverzeichnis", level=1)

        toc_items = [
            ("1. Baustellenbeschrieb", 4),
            ("   1.1 Objektdaten", 4),
            ("   1.2 Gebäudemasse", 4),
            ("   1.3 Gerüstanforderungen", 4),
            ("   1.4 Baustellensituation", 4),
            ("   1.5 Termine", 4),
            ("2. Ausmass", 5),
            ("   2.1 Zeichnungen", 5),
            ("   2.2 Ausmassberechnung nach NPK 114", 5),
            ("   2.3 Vollständiges Ausmass aller Positionen", 6),
            ("   2.4 Bezug auf SIA Norm 118/222", 7),
            ("3. Materialauszug", 8),
            ("   3.1 Materialliste Layher Blitz 70", 8),
            ("   3.2 Gewichtszusammenfassung", 10),
            ("4. Personalbedarf", 11),
            ("5. Dokumentation Baustelle", 12),
            ("   5.1 Materialtransport", 12),
            ("   5.2 Ablad", 12),
            ("   5.3 Umschlagplatz", 12),
            ("   5.4 Sicherheitskonzept", 13),
            ("6. Reflexion", 14),
            ("7. Anhang", 16),
        ]

        for item, page in toc_items:
            p = doc.add_paragraph()
            p.add_run(item)
            # Punkte und Seitenzahl (vereinfacht)

    def _add_section_1(self, doc: Document, building: BuildingData, req: ScaffoldRequirements, project: str):
        """Kapitel 1: Baustellenbeschrieb"""

        doc.add_heading("1. Baustellenbeschrieb", level=1)

        # 1.1 Objektdaten
        doc.add_heading("1.1 Objektdaten", level=2)

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        data = [
            ("Bauvorhaben", project),
            ("Objekt", building.building_category),
            ("Adresse", building.address),
            ("Bauherr", "[Bauherr eintragen]"),
            ("Bauleitung", "[Bauleitung eintragen]"),
            ("Gerüstbauer", f"{self.company_name}, {self.company_address}"),
        ]

        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = value

        # 1.2 Gebäudemasse
        doc.add_heading("1.2 Gebäudemasse", level=2)

        ridge = building.ridge_height_m or (building.eave_height_m + 3.5)
        gable_height = ridge - building.eave_height_m

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        data = [
            ("Grundrissmasse", f"{building.length_m} m × {building.width_m} m (L × B)"),
            ("Traufhöhe", f"{building.eave_height_m} m ab OK Terrain"),
            ("Firsthöhe", f"{ridge} m ab OK Terrain"),
            ("Giebelhöhe", f"{gable_height:.1f} m (über Traufe)"),
            ("Dachform", f"{building.roof_type.capitalize()}, Neigung ca. {building.roof_angle_deg}°"),
            ("Giebel", f"2 Stück an den Schmalseiten ({min(building.length_m, building.width_m)} m)"),
        ]

        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = value

        # 1.3 Gerüstanforderungen
        doc.add_heading("1.3 Gerüstanforderungen", level=2)

        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        data = [
            ("Gerüstart", req.scaffold_type),
            ("Gerüstsystem", req.scaffold_system),
            ("Lastklasse", f"{req.load_class} (200 kg/m²)"),
            ("Breitenklasse", f"{req.width_class} (0.90 m Belagbreite)"),
            ("Fassadenabstand", f"{req.facade_distance_m} m"),
            ("Dachfangschutz", f"Dachfanggerüst an Traufseiten (2 × {max(building.length_m, building.width_m)} m)"),
            ("Giebelgerüst", f"Gerüst über Traufe an Giebelseiten (2 × {min(building.length_m, building.width_m)} m)"),
            ("Verankerung", "Gerüsthalter am Mauerwerk / Fensterrahmen"),
        ]

        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = value

        # 1.4 Baustellensituation
        doc.add_heading("1.4 Baustellensituation", level=2)

        p = doc.add_paragraph()
        p.add_run(f"Das Gebäude befindet sich an der Adresse {building.address}. ")
        p.add_run("Die Zufahrt zur Baustelle ist für LKW bis 18 t möglich. ")
        p.add_run("Das Gelände um das Gebäude ist eben und bietet ausreichend Platz für die Gerüstmontage und Materiallagerung.")

        doc.add_paragraph("• Zufahrt: Strasse, befahrbar für LKW", style='List Bullet')
        doc.add_paragraph("• Terrain: Eben", style='List Bullet')
        doc.add_paragraph("• Nachbarbebauung: [Abstand prüfen]", style='List Bullet')
        doc.add_paragraph("• Hindernisse: [Freileitungen prüfen]", style='List Bullet')
        doc.add_paragraph("• Lagerplatz: Ca. 50 m² verfügbar", style='List Bullet')

        # 1.5 Termine
        doc.add_heading("1.5 Termine", level=2)

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        data = [
            ("Gerüstmontage", "[KW/Jahr] (1 Tag)"),
            ("Vorhaltedauer", "[X] Wochen"),
            ("Gerüstdemontage", "[KW/Jahr] (1 Tag)"),
        ]

        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = value

    def _add_section_2(self, doc: Document, building: BuildingData, npk: NPK114Result):
        """Kapitel 2: Ausmass"""

        doc.add_heading("2. Ausmass", level=1)

        # 2.1 Zeichnungen
        doc.add_heading("2.1 Zeichnungen", level=2)
        p = doc.add_paragraph()
        p.add_run("Die vollständigen Gerüstzeichnungen befinden sich im Anhang A (Grundriss) und Anhang B (Schnitt/Ansicht).")

        # 2.2 Ausmassberechnung
        doc.add_heading("2.2 Ausmassberechnung nach NPK 114", level=2)

        p = doc.add_paragraph()
        p.add_run("Die Ausmassberechnung erfolgt gemäss NPK 114 D/2012 «Arbeitsgerüste» und den Ausmassgrundsätzen im Anhang 1-4.")

        # Grundsätze
        doc.add_heading("2.2.1 Ausmassgrundsätze (NPK 114, Anhang 1)", level=3)

        doc.add_paragraph("• Längen und Höhen: In Meter [m] mit Genauigkeit 0.1 m", style='List Bullet')
        doc.add_paragraph("• Flächen: In Quadratmeter [m²] mit Genauigkeit 0.01 m²", style='List Bullet')
        doc.add_paragraph("• Rundung: Kaufmännisch (0-4 abrunden, 5-9 aufrunden)", style='List Bullet')
        doc.add_paragraph("• Minimale Ausmasslänge: LAmin ≥ 2.5 m", style='List Bullet')
        doc.add_paragraph("• Minimale Ausmasshöhe: HAmin ≥ 4.0 m", style='List Bullet')

        # Zuschläge
        doc.add_heading("2.2.2 Zuschläge (NPK 114, Anhang 2-4)", level=3)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        headers = ["Bezeichnung", "Formelzeichen", "Wert"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        data = [
            ("Fassadenabstand", "LF", "0.30 m"),
            ("Gerüstgangbreite", "LG", "0.70 m (bis 0.70 m)"),
            ("Stirnseitiger Abschluss", "LS", "1.00 m (= LF + LG)"),
            ("Höhenzuschlag", "-", "+ 1.00 m (über Arbeitshöhe)"),
        ]

        for i, row in enumerate(data, 1):
            for j, val in enumerate(row):
                table.cell(i, j).text = val

        # Berechnung
        doc.add_heading("2.2.3 Fassadengerüst – Ausmassberechnung", level=3)

        for calc in npk.calculations:
            p = doc.add_paragraph()
            run = p.add_run(calc["name"])
            run.bold = True

            if "LA_m" in calc:
                p = doc.add_paragraph()
                p.add_run(f"Länge L = {calc['length_m']} m")
                p = doc.add_paragraph()
                p.add_run(f"LA = LS + L + LS = 1.0 + {calc['length_m']} + 1.0 = {calc['LA_m']:.1f} m")
                p = doc.add_paragraph()
                p.add_run(f"Höhe H = {calc['height_m']:.1f} m")
                p = doc.add_paragraph()
                p.add_run(f"HA = H + 1.0 m = {calc['height_m']:.1f} + 1.0 = {calc['HA_m']:.1f} m")

            p = doc.add_paragraph()
            run = p.add_run(f"Fläche: {calc['area_single_m2']:.2f} m² × {calc['count']} = {calc['area_total_m2']:.2f} m²")
            run.bold = True

            doc.add_paragraph()

        # Zusammenfassung
        doc.add_heading("2.2.4 Ausmass-Zusammenfassung", level=3)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        headers = ["Position", "Menge", "Einheit"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        data = [
            ("Fassadengerüst Traufseiten (A+C)", f"{npk.facade_a_c_m2:.2f}", "m²"),
            ("Fassadengerüst Giebelseiten (B+D)", f"{npk.facade_b_d_m2:.2f}", "m²"),
            ("Eckzuschläge (4 Stk.)", f"{npk.corner_surcharge_m2:.2f}", "m²"),
            ("Total Fassadengerüst", f"{npk.facade_area_total_m2:.2f}", "m²"),
        ]

        for i, row in enumerate(data, 1):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = val
                if i == 4:  # Total-Zeile
                    cell.paragraphs[0].runs[0].bold = True

    def _add_section_3(self, doc: Document, material: MaterialList):
        """Kapitel 3: Materialauszug"""

        doc.add_heading("3. Materialauszug", level=1)

        doc.add_heading("3.1 Materialliste Layher Blitz 70", level=2)

        p = doc.add_paragraph()
        p.add_run("Die Materialliste basiert auf dem Gerüstsystem Layher Blitz 70 Stahl für ein umlaufendes Fassadengerüst mit Giebelgerüstung und Dachfangschutz.")

        # Nach Kategorien gruppieren
        categories = {}
        for item in material.items:
            cat = item["category"]
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(item)

        for cat_name, items in categories.items():
            doc.add_heading(f"3.1.x {cat_name}", level=3)

            table = doc.add_table(rows=len(items) + 1, cols=5)
            table.style = 'Table Grid'

            headers = ["Artikel", "Art.-Nr.", "Menge", "kg/Stk", "Total kg"]
            for i, h in enumerate(headers):
                table.cell(0, i).text = h
                table.cell(0, i).paragraphs[0].runs[0].bold = True

            for i, item in enumerate(items, 1):
                table.cell(i, 0).text = item["name"]
                table.cell(i, 1).text = item["art_nr"]
                table.cell(i, 2).text = str(item["quantity"])
                table.cell(i, 3).text = str(item["weight_kg"])
                table.cell(i, 4).text = f"{item['total_kg']:,.0f}".replace(",", "'")

        # Gewichtszusammenfassung
        doc.add_heading("3.2 Gewichtszusammenfassung", level=2)

        table = doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'

        headers = ["Materialgruppe", "Gewicht [kg]"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        data = [
            ("Vertikalrahmen / Stellrahmen", f"{material.vertical_frames_kg:,.0f}".replace(",", "'")),
            ("Horizontalrahmen / Geländer", f"{material.horizontal_frames_kg:,.0f}".replace(",", "'")),
            ("Beläge", f"{material.decks_kg:,.0f}".replace(",", "'")),
            ("Diagonalen und Aussteifung", f"{material.diagonals_kg:,.0f}".replace(",", "'")),
            ("Fussplatten und Spindeln", f"{material.base_plates_kg:,.0f}".replace(",", "'")),
            ("Verankerung", f"{material.anchoring_kg:,.0f}".replace(",", "'")),
            ("Konsolen und Zubehör", f"{material.accessories_kg:,.0f}".replace(",", "'")),
            ("Gesamtgewicht Gerüstmaterial", f"{material.total_weight_kg:,.0f} kg".replace(",", "'")),
        ]

        for i, (label, value) in enumerate(data, 1):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            if i == 8:  # Total-Zeile
                table.cell(i, 0).paragraphs[0].runs[0].bold = True
                table.cell(i, 1).paragraphs[0].runs[0].bold = True

        # Hinweis
        p = doc.add_paragraph()
        p.add_run("Hinweis zum Materialtransport: ").bold = True
        tonnes = material.total_weight_kg / 1000
        p.add_run(f"Das Gesamtgewicht von ca. {tonnes:.1f} Tonnen kann mit einem 3-Achser LKW (Nutzlast ca. 12-14 t) in einer Fuhre transportiert werden.")

    def _add_section_4(self, doc: Document, npk: NPK114Result, personnel: PersonnelRequirement):
        """Kapitel 4: Personalbedarf"""

        doc.add_heading("4. Personalbedarf", level=1)

        # 4.1 Montage
        doc.add_heading("4.1 Montage", level=2)

        p = doc.add_paragraph()
        p.add_run("Für die Montage des umlaufenden Fassadengerüsts inkl. Giebelgerüstung und Dachfangschutz wird folgender Personalbedarf kalkuliert:")

        table = doc.add_table(rows=7, cols=3)
        table.style = 'Table Grid'

        headers = ["Position", "Personal", "Zeit"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        # Zeiten verteilen
        base_time = 1.0
        main_time = personnel.assembly_hours - 2.5

        data = [
            ("Abladen und Bereitstellen", f"{personnel.team_size} Gerüstbauer", "1.0 h"),
            (f"Montage Fassadengerüst (ca. {npk.facade_area_total_m2:.0f} m²)", f"{personnel.team_size} Gerüstbauer", f"{main_time:.1f} h"),
            ("Montage Dachfangschutz", f"{personnel.team_size} Gerüstbauer", "1.0 h"),
            ("Montage Giebelüberstand", f"{personnel.team_size} Gerüstbauer", "1.0 h"),
            ("Verankerung und Kontrolle", f"{personnel.team_size} Gerüstbauer", "0.5 h"),
            ("Total Montage", f"{personnel.team_size} Gerüstbauer", f"{personnel.assembly_hours:.1f} h (1 Tag)"),
        ]

        for i, row in enumerate(data, 1):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = val
                if i == 6:
                    cell.paragraphs[0].runs[0].bold = True

        # 4.2 Demontage
        doc.add_heading("4.2 Demontage", level=2)

        p = doc.add_paragraph()
        p.add_run("Die Demontage erfolgt in umgekehrter Reihenfolge. Erfahrungsgemäss ist die Demontage ca. 20% schneller als die Montage.")

        p = doc.add_paragraph()
        run = p.add_run(f"Total Demontage: {personnel.team_size} Gerüstbauer, {personnel.disassembly_hours:.1f} h (1 Tag)")
        run.bold = True

        # 4.3 Zusammenfassung
        doc.add_heading("4.3 Zusammenfassung Personalbedarf", level=2)

        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'

        headers = ["Arbeitsphase", "Personal", "Dauer", "Mannstunden"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        mh_montage = personnel.assembly_hours * personnel.team_size
        mh_demontage = personnel.disassembly_hours * personnel.team_size

        data = [
            ("Montage", f"{personnel.team_size} Pers.", f"{personnel.assembly_hours:.1f} h", f"{mh_montage:.1f} Mh"),
            ("Demontage", f"{personnel.team_size} Pers.", f"{personnel.disassembly_hours:.1f} h", f"{mh_demontage:.1f} Mh"),
            ("Total", f"{personnel.team_size} Pers.", f"{personnel.assembly_hours + personnel.disassembly_hours:.1f} h", f"{personnel.total_man_hours:.1f} Mh"),
        ]

        for i, row in enumerate(data, 1):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = val
                if i == 3:
                    cell.paragraphs[0].runs[0].bold = True

    def _add_section_5(self, doc: Document, building: BuildingData, material: MaterialList):
        """Kapitel 5: Dokumentation Baustelle"""

        doc.add_heading("5. Dokumentation Baustelle", level=1)

        # 5.1 Materialtransport
        doc.add_heading("5.1 Materialtransport", level=2)

        p = doc.add_paragraph()
        run = p.add_run("Transportmittel:")
        run.bold = True
        doc.add_paragraph("3-Achs-LKW mit Pritsche und Kran (HIAB), Nutzlast ca. 12-14 Tonnen")

        p = doc.add_paragraph()
        run = p.add_run("Ladungssicherung:")
        run.bold = True
        doc.add_paragraph("• Stellrahmen gebündelt und mit Spanngurten gesichert", style='List Bullet')
        doc.add_paragraph("• Beläge in Gitterboxen oder auf Paletten gestapelt", style='List Bullet')
        doc.add_paragraph("• Kleinmaterial in beschrifteten Kisten", style='List Bullet')

        # 5.2 Ablad
        doc.add_heading("5.2 Ablad", level=2)

        p = doc.add_paragraph()
        run = p.add_run("Abladeort:")
        run.bold = True
        doc.add_paragraph(f"Bei {building.address}")

        # 5.3 Umschlagplatz
        doc.add_heading("5.3 Umschlagplatz (Platzbedarf)", level=2)

        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        headers = ["Verwendung", "Fläche", "Bemerkung"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        data = [
            ("Materiallager Stellrahmen", "15 m²", "Gestapelt, max. 1.5 m hoch"),
            ("Materiallager Beläge", "12 m²", "Paletten / Gitterbox"),
            ("Materiallager Geländer", "10 m²", "Gebündelt, liegend"),
            ("Kleinmaterial und Zubehör", "5 m²", "In Kisten"),
            ("Total Platzbedarf", "ca. 50 m²", ""),
        ]

        for i, row in enumerate(data, 1):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = val
                if i == 5:
                    cell.paragraphs[0].runs[0].bold = True

        # 5.4 Sicherheitskonzept
        doc.add_heading("5.4 Sicherheitskonzept", level=2)

        p = doc.add_paragraph()
        p.add_run("Das Sicherheitskonzept basiert auf der Gefährdungsbeurteilung gemäss BauAV und den Vorgaben der SUVA.")

        doc.add_heading("5.4.1 Gefährdungsbeurteilung", level=3)

        table = doc.add_table(rows=7, cols=3)
        table.style = 'Table Grid'

        headers = ["Gefährdung", "Massnahme", "Verantwortlich"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].bold = True

        data = [
            ("Absturz während Montage", "MSG-Verfahren, PSAgA ab 3. Lage", "Gruppenleiter"),
            ("Herabfallende Teile", "Helm, Absperrung Gefahrenbereich", "Alle Mitarbeiter"),
            ("Stolpern/Ausrutschen", "Ordnung, Sicherheitsschuhe S3", "Alle Mitarbeiter"),
            ("Manuelle Lasthandhabung", "Hebebilfen, max. 25 kg/Person", "Gruppenleiter"),
            ("Verkehr während Ablad", "Absperrung, Signalisation", "Fahrer / GL"),
            ("Witterung", "Arbeitsunterbruch bei Sturm/Gewitter", "Gruppenleiter"),
        ]

        for i, row in enumerate(data, 1):
            for j, val in enumerate(row):
                table.cell(i, j).text = val

        doc.add_heading("5.4.2 Persönliche Schutzausrüstung (PSA)", level=3)
        doc.add_paragraph("• Schutzhelm (obligatorisch)", style='List Bullet')
        doc.add_paragraph("• Sicherheitsschuhe S3", style='List Bullet')
        doc.add_paragraph("• Arbeitshandschuhe", style='List Bullet')
        doc.add_paragraph("• PSAgA (Auffanggurt, Verbindungsmittel) für Montage", style='List Bullet')
        doc.add_paragraph("• Signalweste bei Arbeiten im Verkehrsbereich", style='List Bullet')

    def _add_section_6(self, doc: Document):
        """Kapitel 6: Reflexion (Vorlage zum Ausfüllen)"""

        doc.add_heading("6. Reflexion", level=1)

        # 6.1 Planungsphase
        doc.add_heading("6.1 Planungsphase", level=2)

        p = doc.add_paragraph()
        run = p.add_run("Leitfragen:")
        run.italic = True

        doc.add_paragraph("• Was lief gut bei der Materialdisposition?", style='List Bullet')
        doc.add_paragraph("• Welche Herausforderungen gab es bei der Mengenermittlung?", style='List Bullet')
        doc.add_paragraph("• Wie wurde die Kommunikation mit dem Lager / der Disposition geführt?", style='List Bullet')

        p = doc.add_paragraph()
        p.add_run("[Hier Ihre Reflexion zur Planungsphase eintragen...]")

        # 6.2 Ausführungsphase
        doc.add_heading("6.2 Ausführungsphase", level=2)

        p = doc.add_paragraph()
        run = p.add_run("Leitfragen:")
        run.italic = True

        doc.add_paragraph("• War das bestellte Material vollständig und korrekt?", style='List Bullet')
        doc.add_paragraph("• Wie verlief der Transport und Ablad?", style='List Bullet')
        doc.add_paragraph("• Gab es Material-Engpässe oder Überschüsse?", style='List Bullet')

        p = doc.add_paragraph()
        p.add_run("[Hier Ihre Reflexion zur Ausführungsphase eintragen...]")

        # 6.3 Erkenntnisse
        doc.add_heading("6.3 Erkenntnisse und Verbesserungspotential", level=2)

        p = doc.add_paragraph()
        p.add_run("[Hier Ihre Erkenntnisse eintragen...]")

        # 6.4 Persönliches Fazit
        doc.add_heading("6.4 Persönliches Fazit", level=2)

        p = doc.add_paragraph()
        p.add_run("[Hier Ihr persönliches Fazit eintragen...]")

    def _add_section_7(self, doc: Document, building: BuildingData):
        """Kapitel 7: Anhang"""

        doc.add_heading("7. Anhang", level=1)

        doc.add_heading("Anhang A: Grundriss Gerüst", level=2)
        p = doc.add_paragraph()
        p.add_run("[SVG-Visualisierung wird hier eingefügt - siehe API-Endpunkt /api/v1/visualize/floor-plan]")

        doc.add_heading("Anhang B: Schnitt / Ansicht Giebelseite", level=2)
        p = doc.add_paragraph()
        p.add_run("[SVG-Visualisierung wird hier eingefügt - siehe API-Endpunkt /api/v1/visualize/cross-section]")

        doc.add_heading("Anhang C: Gerüstkarte / Kennzeichnung", level=2)

        # Gerüstkarte-Vorlage
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Table Grid'

        data = [
            ("Gerüstersteller (Firma):", f"{self.company_name}"),
            ("Baustelle / Objekt:", building.address),
            ("Gerüstart:", "Fassadengerüst"),
            ("Gerüstsystem:", "Layher Blitz 70 Stahl"),
            ("Lastklasse:", "3 (200 kg/m²)"),
            ("Breitenklasse:", "W09 (0.90 m)"),
            ("Gerüsthöhe:", f"{building.eave_height_m + 2:.1f} m"),
            ("Gerüstlänge total:", f"ca. {2 * (building.length_m + building.width_m):.0f} m (umlaufend)"),
            ("Datum der Freigabe:", "_________________"),
            ("Unterschrift Prüfer:", "_________________"),
        ]

        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = value

        doc.add_heading("Anhang D: Checkliste Materialkontrolle", level=2)
        p = doc.add_paragraph()
        p.add_run("[Checkliste basierend auf Materialliste - siehe Kapitel 3]")


# Singleton-Instanz
_generator = None

def get_document_generator() -> DocumentGenerator:
    """Gibt die Singleton-Instanz des DocumentGenerators zurück"""
    global _generator
    if _generator is None:
        _generator = DocumentGenerator()
    return _generator
